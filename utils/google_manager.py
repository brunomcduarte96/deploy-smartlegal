import json
from typing import List, Dict, Any, Tuple
from google.oauth2 import service_account
from googleapiclient.discovery import build
from googleapiclient.http import MediaIoBaseUpload, MediaFileUpload
from io import BytesIO
from config.settings import (
    GOOGLE_CREDENTIALS,
    SHEETS_SCOPE,
    DRIVE_SCOPE,
    DOCS_SCOPE,
    SHEET_ID_1,
    SHEET_ID_2,
    ROOT_FOLDER_ID
)
from utils.date_utils import data_por_extenso
from utils.error_handler import DriveError
from datetime import datetime
from docx import Document
import re
import os
from pathlib import Path
import logging
from unidecode import unidecode
import pytz
import tempfile
import time

# Configurar timezone de São Paulo
SP_TZ = pytz.timezone('America/Sao_Paulo')

logger = logging.getLogger(__name__)

logger.info(f"Diretório atual: {os.getcwd()}")

class GoogleManager:
    def __init__(self):
        self.credentials = self._get_credentials()
        self.sheets_service = self._build_sheets_service()
        self.drive_service = self._build_drive_service()
        self.docs_service = self._build_docs_service()

    def _get_credentials(self):
        """Cria credenciais do Google a partir do JSON armazenado"""
        creds_dict = json.loads(GOOGLE_CREDENTIALS)
        scopes = SHEETS_SCOPE + DRIVE_SCOPE + DOCS_SCOPE
        return service_account.Credentials.from_service_account_info(
            creds_dict, scopes=scopes
        )

    def _build_sheets_service(self):
        """Cria serviço do Google Sheets"""
        return build('sheets', 'v4', credentials=self.credentials)

    def _build_drive_service(self):
        """Cria serviço do Google Drive"""
        return build('drive', 'v3', credentials=self.credentials)

    def _build_docs_service(self):
        """Cria serviço do Google Docs"""
        return build('docs', 'v1', credentials=self.credentials)

    def update_sheet(self, sheet_id: str, range_name: str, values: List[List[Any]]):
        """Adiciona dados na última linha da planilha do Google Sheets"""
        try:
            # Primeiro, encontra a última linha com dados
            result = self.sheets_service.spreadsheets().values().get(
                spreadsheetId=sheet_id,
                range=range_name
            ).execute()
            
            # Pega o número da próxima linha vazia
            last_row = len(result.get('values', [])) + 1
            
            # Ajusta o range para apontar para a próxima linha vazia
            column_range = range_name.split(':')[1]
            new_range = f"A{last_row}:{column_range}{last_row}"
            
            # Faz o append dos dados
            body = {'values': values}
            self.sheets_service.spreadsheets().values().append(
                spreadsheetId=sheet_id,
                range=new_range,
                valueInputOption='USER_ENTERED',
                insertDataOption='INSERT_ROWS',
                body=body
            ).execute()
            
            logger.info(f"Dados adicionados na linha {last_row} da planilha {sheet_id}")
            
        except Exception as e:
            logger.error(f"Erro ao atualizar planilha: {str(e)}")
            raise Exception(f"Erro ao atualizar planilha: {str(e)}")

    def create_folder(self, folder_name: str, parent_id: str = ROOT_FOLDER_ID) -> str:
        """Cria uma pasta no Google Drive"""
        try:
            file_metadata = {
                'name': folder_name,
                'mimeType': 'application/vnd.google-apps.folder',
                'parents': [parent_id]
            }
            file = self.drive_service.files().create(
                body=file_metadata,
                fields='id'
            ).execute()
            return file.get('id')
        except Exception as e:
            raise Exception(f"Erro ao criar pasta: {str(e)}")

    def upload_file(self, file_name: str, file_content: bytes, mime_type: str, folder_id: str) -> str:
        """Faz upload de um arquivo para o Google Drive"""
        try:
            file_metadata = {
                'name': file_name,
                'parents': [folder_id]
            }
            media = MediaIoBaseUpload(
                BytesIO(file_content),
                mimetype=mime_type,
                resumable=True
            )
            file = self.drive_service.files().create(
                body=file_metadata,
                media_body=media,
                fields='id'
            ).execute()
            return file.get('id')
        except Exception as e:
            raise Exception(f"Erro ao fazer upload do arquivo: {str(e)}")

    def fill_document_template(self, template_path: str, data: Dict[str, str], folder_id: str, output_filename: str = None) -> Tuple[str, str]:
        """
        Preenche o template e salva como PDF e DOCX
        
        Args:
            template_path: Caminho para o template
            data: Dicionário com os dados para substituição
            folder_id: ID da pasta onde salvar os arquivos
            output_filename: Nome personalizado para o arquivo de saída (sem extensão)
            
        Retorna: (pdf_id, docx_id)
        """
        try:
            # Verifica se o arquivo existe (usando caminho absoluto)
            template_full_path = os.path.join(os.getcwd(), 'templates', template_path)
            if not os.path.exists(template_full_path):
                raise DriveError(f"Template não encontrado: {template_full_path}")
            
            logger.info(f"Usando template em: {template_full_path}")

            # Carrega o template
            try:
                doc = Document(template_full_path)
            except Exception as e:
                raise DriveError(f"Erro ao carregar template: {str(e)}")
            
            # Log dos dados
            logger.info(f"Dados para substituição: {data}")
            
            # Substitui os placeholders em todo o documento
            for paragraph in doc.paragraphs:
                original_text = paragraph.text
                for key, value in data.items():
                    if f"{{{{{key}}}}}" in paragraph.text:
                        paragraph.text = paragraph.text.replace(f"{{{{{key}}}}}", str(value))
                
                # Log se houve substituição
                if original_text != paragraph.text:
                    logger.debug(f"Substituído: '{original_text}' -> '{paragraph.text}'")
            
            # Substitui os placeholders nas tabelas
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            original_text = paragraph.text
                            for key, value in data.items():
                                if f"{{{{{key}}}}}" in paragraph.text:
                                    paragraph.text = paragraph.text.replace(f"{{{{{key}}}}}", str(value))
                            
                            # Log se houve substituição
                            if original_text != paragraph.text:
                                logger.debug(f"Substituído em tabela: '{original_text}' -> '{paragraph.text}'")
            
            # Define o nome do arquivo temporário
            temp_docx_path = os.path.join(tempfile.gettempdir(), f"temp_{int(time.time())}.docx")
            
            # Salva o documento temporário
            doc.save(temp_docx_path)
            logger.info(f"Documento temporário salvo em: {temp_docx_path}")
            
            # Define o nome do arquivo final
            if output_filename:
                file_name = output_filename
            else:
                # Usa o nome do template sem a extensão
                file_name = os.path.splitext(os.path.basename(template_path))[0]
            
            # Upload do DOCX para o Drive
            docx_metadata = {
                'name': f"{file_name}.docx",
                'parents': [folder_id],
                'mimeType': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            }
            
            docx_media = MediaFileUpload(temp_docx_path, mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
            docx_file = self.drive_service.files().create(body=docx_metadata, media_body=docx_media, fields='id').execute()
            docx_id = docx_file.get('id')
            
            logger.info(f"DOCX enviado para o Drive. ID: {docx_id}")
            
            # Converte para PDF
            pdf_metadata = {
                'name': f"{file_name}.pdf",
                'parents': [folder_id],
                'mimeType': 'application/pdf'
            }
            
            self.drive_service.files().copy(
                fileId=docx_id,
                body=pdf_metadata,
                fields='id'
            ).execute()
            
            # Busca o PDF recém-criado
            pdf_files = self.drive_service.files().list(
                q=f"name='{file_name}.pdf' and '{folder_id}' in parents",
                spaces='drive',
                fields='files(id, name)'
            ).execute().get('files', [])
            
            if not pdf_files:
                raise DriveError("PDF não encontrado após conversão")
            
            pdf_id = pdf_files[0].get('id')
            logger.info(f"PDF criado no Drive. ID: {pdf_id}")
            
            # Limpa o arquivo temporário
            try:
                os.remove(temp_docx_path)
                logger.info(f"Arquivo temporário removido: {temp_docx_path}")
            except Exception as e:
                logger.warning(f"Erro ao remover documento temporário: {str(e)}")
            
            return pdf_id, docx_id
            
        except Exception as e:
            logger.error(f"Erro ao processar template: {str(e)}")
            raise DriveError(f"Erro ao processar template: {str(e)}")

    def export_to_pdf(self, doc_id: str) -> bytes:
        """Exporta um documento do Google Docs para PDF"""
        try:
            return self.drive_service.files().export(
                fileId=doc_id,
                mimeType='application/pdf'
            ).execute()
        except Exception as e:
            raise Exception(f"Erro ao exportar para PDF: {str(e)}")

    def update_sheets_with_client_data(self, client_data: Dict[str, Any], folder_url: str, caso_data: Dict[str, Any], is_new_client: bool = False):
        """
        Atualiza as duas planilhas com os dados do cliente e do caso
        
        Args:
            client_data: Dados do cliente
            folder_url: URL da pasta do cliente
            caso_data: Dados do caso
            is_new_client: Se True, adiciona cliente na primeira planilha
        """
        try:
            # Atualiza primeira planilha apenas se for cliente novo
            if is_new_client:
                # Formata a data de nascimento
                data_nascimento = datetime.fromisoformat(client_data['data_nascimento'])
                data_nascimento_formatada = data_nascimento.strftime('%d/%m/%Y')
                
                # Atualiza primeira planilha (dados pessoais do cliente)
                values1 = [[
                    client_data['nome_completo'],
                    client_data['nacionalidade'],
                    client_data['estado_civil'],
                    client_data['profissao'],
                    client_data['email'],
                    client_data['celular'],
                    data_nascimento_formatada,
                    client_data['rg'],
                    client_data['cpf'],
                    client_data['endereco'],
                    client_data['bairro'],
                    client_data['cidade'],
                    client_data['estado'],
                    client_data['cep'],
                    folder_url  # URL da pasta do cliente
                ]]
                
                # Append direto na primeira planilha
                self.sheets_service.spreadsheets().values().append(
                    spreadsheetId=SHEET_ID_1,
                    range='A:O',
                    valueInputOption='USER_ENTERED',
                    insertDataOption='INSERT_ROWS',
                    body={'values': values1}
                ).execute()
                logger.info(f"Planilha 1 atualizada com dados de {client_data['nome_completo']}")
            
            # Sempre atualiza a segunda planilha com o novo caso
            current_date = datetime.now(SP_TZ).strftime('%d/%m/%Y')
            values2 = [[
                client_data['nome_completo'],          # A: Cliente
                caso_data['caso'],                     # B: Caso
                caso_data['assunto_caso'],             # C: Assunto Caso
                "Em andamento",                        # D: Status
                "",                                    # E: Vazio
                client_data['estado'],                 # F: Estado
                caso_data['pasta_caso_url'],           # G: Pasta Drive
                current_date,                          # H: Data Entrada
                "SMARTLEGAL",                          # I: SMARTLEGAL
                caso_data['responsavel_comercial']     # J: Responsavel Comercial
            ]]
            
            # Append direto na segunda planilha
            self.sheets_service.spreadsheets().values().append(
                spreadsheetId=SHEET_ID_2,
                range='A:J',  # Atualizado para incluir até a coluna J
                valueInputOption='USER_ENTERED',
                insertDataOption='INSERT_ROWS',
                body={'values': values2}
            ).execute()
            logger.info(f"Planilha 2 atualizada com caso para {client_data['nome_completo']}")
            
        except Exception as e:
            logger.error(f"Erro ao atualizar planilhas: {str(e)}")
            raise Exception(f"Erro ao atualizar planilhas: {str(e)}")

    def format_folder_name(self, text: str) -> str:
        """Formata nome da pasta removendo caracteres especiais e espaços"""
        text = text.strip()
        text = unidecode(text)
        text = text.replace(' ', '_')
        return text

    def get_or_create_client_folder(self, nome: str, cpf: str) -> str:
        """
        Busca ou cria pasta do cliente no formato NOME_CPF
        Retorna: ID da pasta do cliente
        """
        try:
            nome_formatado = self.format_folder_name(nome)
            cpf_formatado = cpf.replace('.', '').replace('-', '')
            folder_name = f"{nome_formatado}_{cpf_formatado}"
            
            # Verifica se pasta já existe
            response = self.drive_service.files().list(
                q=f"name='{folder_name}' and '{ROOT_FOLDER_ID}' in parents and mimeType='application/vnd.google-apps.folder'",
                spaces='drive'
            ).execute()
            
            if response.get('files'):
                logger.info(f"Pasta do cliente encontrada: {folder_name}")
                return response['files'][0]['id']
            
            # Cria nova pasta
            folder_metadata = {
                'name': folder_name,
                'mimeType': 'application/vnd.google-apps.folder',
                'parents': [ROOT_FOLDER_ID]
            }
            folder = self.drive_service.files().create(
                body=folder_metadata,
                fields='id'
            ).execute()
            
            logger.info(f"Pasta do cliente criada: {folder_name}")
            return folder.get('id')
            
        except Exception as e:
            raise DriveError(f"Erro ao criar/buscar pasta do cliente: {str(e)}")

    def create_case_folder(self, client_folder_id: str, assunto_caso: str) -> str:
        """
        Cria pasta do caso no formato ASSUNTO_CASO_DATA
        Retorna: ID da pasta do caso
        """
        try:
            assunto_formatado = self.format_folder_name(assunto_caso)
            data_atual = datetime.now(SP_TZ).strftime('%Y%m%d')
            folder_name = f"{assunto_formatado}_{data_atual}"
            
            folder_metadata = {
                'name': folder_name,
                'mimeType': 'application/vnd.google-apps.folder',
                'parents': [client_folder_id]
            }
            
            folder = self.drive_service.files().create(
                body=folder_metadata,
                fields='id'
            ).execute()
            
            logger.info(f"Pasta do caso criada: {folder_name}")
            return folder.get('id')
            
        except Exception as e:
            raise DriveError(f"Erro ao criar pasta do caso: {str(e)}")

    def get_folder_url(self, folder_id: str) -> str:
        """Retorna URL da pasta do Drive"""
        try:
            # Verifica se a pasta existe
            self.drive_service.files().get(fileId=folder_id).execute()
            return f"https://drive.google.com/drive/folders/{folder_id}"
        except Exception as e:
            raise DriveError(f"Erro ao gerar URL da pasta: {str(e)}")

    def get_files_in_folder(self, folder_id: str, query: str = None):
        """
        Busca arquivos em uma pasta do Google Drive
        
        Args:
            folder_id: ID da pasta
            query: Query adicional para filtrar arquivos
        
        Returns:
            Lista de arquivos
        """
        try:
            q = f"'{folder_id}' in parents"
            if query:
                q += f" and {query}"
            
            response = self.drive_service.files().list(
                q=q,
                spaces='drive',
                fields='files(id, name, mimeType, webViewLink)'
            ).execute()
            
            return response.get('files', [])
        except Exception as e:
            logger.error(f"Erro ao buscar arquivos na pasta: {str(e)}")
            raise Exception(f"Erro ao buscar arquivos na pasta: {str(e)}") 